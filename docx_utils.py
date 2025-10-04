# docx_utils.py
import os
import subprocess
import tempfile
from pathlib import Path
from docx import Document

def replace_placeholders_in_paragraph(paragraph, placeholders):
    # gather full text
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text
    for k, v in placeholders.items():
        new_text = new_text.replace(k, v)
    if new_text != full_text:
        # Put the whole new_text in the first run and clear others (preserves some styling)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for r in paragraph.runs[1:]:
                r.text = ""
        else:
            paragraph.add_run(new_text)


def replace_placeholders(doc: Document, placeholders: dict):
    # paragraphs
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph(p, placeholders)
    # tables (cells have paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, placeholders)


def load_and_replace(template_path: str, output_path: str, placeholders: dict):
    doc = Document(template_path)
    replace_placeholders(doc, placeholders)
    doc.save(output_path)


def docx_to_pdf(docx_path: str) -> str:
    """
    DOCX -> PDF converter.
    Returns path to converted PDF.
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Input DOCX not found: {docx_path}")

    docx_path = os.path.abspath(docx_path)
    output_dir = tempfile.gettempdir()
    pdf_path = os.path.join(output_dir, Path(docx_path).stem + ".pdf")

    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        if not os.path.exists(pdf_path):
            raise RuntimeError("LibreOffice did not produce output")
        return pdf_path
    except FileNotFoundError:
        raise RuntimeError("LibreOffice not installed. Install with: sudo apt install libreoffice")
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"LibreOffice conversion failed: {e.stderr.decode()}")