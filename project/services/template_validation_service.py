from __future__ import annotations

import datetime as dt
import shutil
import tempfile
from pathlib import Path
from typing import Any, Dict, List

from docx import Document

from project.core import config
from project.services.settings_service import set_setting
from project.utils.docs.docx_replace_placeholders import replace_dynamic_text
from project.utils.docs.pdf_converter import convert_docx_to_pdf
from project.utils.logging_utils import log_error

REQUIRED_PLACEHOLDERS = [
    "{{DJELOVODNI_BROJ}}",
    "{{DANASNJI_DATUM}}",
    "{{IME}}",
    "{{RODITELJ}}",
    "{{DATUM_RODJENJA}}",
    "{{MJESTO}}",
    "{{OPSTINA}}",
    "{{RAZRED}}",
    "{{STRUKA}}",
    "{{RAZLOG}}",
]

SAMPLE_PLACEHOLDERS = {
    "{{DJELOVODNI_BROJ}}": "01-999/26",
    "{{DANASNJI_DATUM}}": "15.03.2026",
    "{{IME}}": "Пробни Ученик",
    "{{RODITELJ}}": "Пробни Родитељ",
    "{{DATUM_RODJENJA}}": "01.01.2008",
    "{{MJESTO}}": "Касиндо",
    "{{OPSTINA}}": "Источна Илиџа",
    "{{RAZRED}}": "ДРУГИ",
    "{{STRUKA}}": "ЕЛЕКТРОТЕХНИКА",
    "{{RAZLOG}}": "ПОТВРДА О СТАТУСУ",
}

VALIDATION_DIR = config.VAR_DIR / "template_validation"


def _collect_text(doc: Document) -> str:
    chunks: List[str] = []
    for para in doc.paragraphs:
        chunks.append(para.text or "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    chunks.append(para.text or "")
    return "\n".join(chunks)


def extract_placeholders(template_path: str) -> Dict[str, Any]:
    path = Path(template_path)
    doc = Document(str(path))
    full_text = _collect_text(doc)
    found = sorted({token for token in REQUIRED_PLACEHOLDERS if token in full_text})
    missing = [token for token in REQUIRED_PLACEHOLDERS if token not in found]
    return {
        "template_path": str(path),
        "found": found,
        "missing_required": missing,
        "text_length": len(full_text),
    }


def validate_template_file(template_path: str, *, run_probe_render: bool = True) -> Dict[str, Any]:
    path = Path(template_path)
    report: Dict[str, Any] = {
        "ok": False,
        "template_path": str(path),
        "file_exists": path.exists(),
        "fatal_issues": [],
        "warnings": [],
        "found_placeholders": [],
        "missing_required": [],
        "probe_docx_path": "",
        "probe_pdf_path": "",
        "probe_render_ok": False,
        "validated_at": dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    if not path.exists() or not path.is_file():
        report["fatal_issues"].append("Template fajl ne postoji ili nije regularan fajl.")
        return report
    if path.suffix.lower() != ".docx":
        report["fatal_issues"].append("Template mora biti .docx fajl.")
        return report

    try:
        doc = Document(str(path))
        full_text = _collect_text(doc)
        found = sorted({token for token in REQUIRED_PLACEHOLDERS if token in full_text})
        missing = [token for token in REQUIRED_PLACEHOLDERS if token not in found]
        report["found_placeholders"] = found
        report["missing_required"] = missing
        if missing:
            report["fatal_issues"].append(
                "Nedostaju obavezni placeholderi: " + ", ".join(missing)
            )
        if len(found) == len(REQUIRED_PLACEHOLDERS):
            report["warnings"].append("Svi obavezni placeholderi su pronađeni.")
    except Exception as e:
        report["fatal_issues"].append(f"DOCX nije moguće otvoriti: {e}")
        return report

    if run_probe_render:
        VALIDATION_DIR.mkdir(parents=True, exist_ok=True)
        try:
            with tempfile.TemporaryDirectory(dir=str(VALIDATION_DIR), prefix="probe_") as tmp:
                tmp_path = Path(tmp)
                probe_docx = tmp_path / "probe_output.docx"
                replace_dynamic_text(str(path), str(probe_docx), SAMPLE_PLACEHOLDERS)
                if not probe_docx.exists() or probe_docx.stat().st_size == 0:
                    raise RuntimeError("Probe DOCX je prazan ili nije kreiran.")
                report["probe_docx_path"] = str(probe_docx)
                try:
                    probe_pdf = Path(convert_docx_to_pdf(str(probe_docx), output_dir=str(tmp_path)))
                    if not probe_pdf.exists() or probe_pdf.stat().st_size == 0:
                        raise RuntimeError("Probe PDF je prazan ili nije kreiran.")
                    report["probe_pdf_path"] = str(probe_pdf)
                    report["probe_render_ok"] = True
                except Exception as e:
                    report["warnings"].append(f"Probe PDF render nije prošao: {e}")
        except Exception as e:
            report["fatal_issues"].append(f"Probe DOCX render nije uspio: {e}")

    report["ok"] = not report["fatal_issues"]
    return report


def remember_template_validation(report: Dict[str, Any]) -> None:
    try:
        set_setting("template_last_validated_at", str(report.get("validated_at") or ""))
        set_setting("template_last_validation_ok", "1" if report.get("ok") else "0")
        set_setting("template_last_validation_summary", summarize_validation_report(report))
        set_setting("template_last_validation_path", str(report.get("template_path") or ""))
    except Exception as e:
        log_error(f"[TEMPLATE] Failed to persist validation state: {e}")


def summarize_validation_report(report: Dict[str, Any]) -> str:
    if report.get("ok"):
        found = len(report.get("found_placeholders") or [])
        return f"Template validan. Placeholders: {found}/{len(REQUIRED_PLACEHOLDERS)}. Probe render: {'OK' if report.get('probe_render_ok') else 'NE'}"
    fatals = report.get("fatal_issues") or []
    return fatals[0] if fatals else "Template validacija nije prošla."


def format_validation_report(report: Dict[str, Any]) -> str:
    lines = [
        f"Template: {report.get('template_path', '-')}",
        f"Validirano: {report.get('validated_at', '-')}",
        f"Status: {'OK' if report.get('ok') else 'FAIL'}",
        f"Probe render: {'OK' if report.get('probe_render_ok') else 'WARN/FAIL'}",
        "",
        f"Pronađeni placeholderi ({len(report.get('found_placeholders') or [])}/{len(REQUIRED_PLACEHOLDERS)}):",
    ]
    lines.extend(f"- {token}" for token in (report.get("found_placeholders") or []))
    missing = report.get("missing_required") or []
    lines.append("")
    lines.append("Nedostajući placeholderi:" if missing else "Nedostajući placeholderi: nema")
    if missing:
        lines.extend(f"- {token}" for token in missing)
    warnings = report.get("warnings") or []
    if warnings:
        lines.append("")
        lines.append("Napomene:")
        lines.extend(f"- {row}" for row in warnings)
    fatals = report.get("fatal_issues") or []
    if fatals:
        lines.append("")
        lines.append("Greške:")
        lines.extend(f"- {row}" for row in fatals)
    return "\n".join(lines)
