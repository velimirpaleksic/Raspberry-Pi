# utils/docs/docx_replace_placeholders.py
from __future__ import annotations

from collections.abc import Iterable
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from typing import Dict


DEFAULT_FONT_SIZE_PT = 11.0
MIN_FIT_FONT_SIZE_PT = 7.0
TEXT_WIDTH_FACTOR = 0.54


def replace_dynamic_text(template_path: str, output_path: str, placeholders: Dict[str, str]):
    """Replace placeholders in DOCX template and save to output."""
    doc = Document(template_path)
    if placeholders:
        _replace_placeholders(doc, placeholders)
    doc.save(output_path)


def _iter_all_paragraphs(doc: Document) -> Iterable:
    for para in doc.paragraphs:
        yield para

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _replace_in_single_run(paragraph, key: str, value: str) -> bool:
    changed = False
    for run in paragraph.runs:
        if key in run.text:
            run.text = run.text.replace(key, value)
            _fit_replacement_runs(paragraph, [run], value)
            changed = True
    return changed


def _replace_across_runs(paragraph, key: str, value: str) -> bool:
    runs = paragraph.runs
    if not runs:
        return False

    full_text = "".join(run.text for run in runs)
    start = full_text.find(key)
    if start < 0:
        return False

    end = start + len(key)
    cursor = 0
    segments: list[tuple[int, int, int]] = []
    for idx, run in enumerate(runs):
        run_text = run.text
        next_cursor = cursor + len(run_text)
        segments.append((idx, cursor, next_cursor))
        cursor = next_cursor

    replacement_start_run = None
    replacement_end_run = None
    for idx, seg_start, seg_end in segments:
        if replacement_start_run is None and start < seg_end:
            replacement_start_run = idx
        if replacement_start_run is not None and end <= seg_end:
            replacement_end_run = idx
            break

    if replacement_start_run is None or replacement_end_run is None:
        return False

    start_run_start = segments[replacement_start_run][1]
    end_run_start = segments[replacement_end_run][1]

    prefix = runs[replacement_start_run].text[: max(0, start - start_run_start)]
    suffix = runs[replacement_end_run].text[max(0, end - end_run_start):]

    if replacement_end_run == replacement_start_run:
        runs[replacement_start_run].text = prefix + value + suffix
        _fit_replacement_runs(paragraph, [runs[replacement_start_run]], value)
        return True

    runs[replacement_start_run].text = prefix + value
    for idx in range(replacement_start_run + 1, replacement_end_run):
        runs[idx].text = ""
    runs[replacement_end_run].text = suffix
    _fit_replacement_runs(paragraph, [runs[replacement_start_run]], value)
    return True


def _find_ancestor(element, tag_name: str):
    current = element
    while current is not None:
        if current.tag == tag_name:
            return current
        current = current.getparent()
    return None


def _cell_width_points(paragraph) -> float | None:
    cell = _find_ancestor(paragraph._p, qn("w:tc"))
    if cell is None:
        return None

    for element in cell.iter():
        if element.tag != qn("w:tcW"):
            continue
        width = element.get(qn("w:w"))
        width_type = element.get(qn("w:type"))
        if not width or width_type not in (None, "dxa"):
            continue
        try:
            width_twips = int(width)
        except ValueError:
            continue
        if width_twips > 0:
            # A little padding keeps text away from cell borders and prevents near-edge clipping.
            return max(48.0, (width_twips / 20.0) - 14.0)
    return None


def _available_width_points(paragraph) -> float:
    cell_width = _cell_width_points(paragraph)
    if cell_width:
        return cell_width

    try:
        section = paragraph.part.document.sections[0]
        return max(96.0, (section.page_width - section.left_margin - section.right_margin).pt)
    except Exception:
        return 450.0


def _base_font_size_points(run, paragraph) -> float:
    try:
        if run.font.size is not None:
            return float(run.font.size.pt)
    except Exception:
        pass

    try:
        if paragraph.style is not None and paragraph.style.font.size is not None:
            return float(paragraph.style.font.size.pt)
    except Exception:
        pass

    return DEFAULT_FONT_SIZE_PT


def _weighted_character_count(text: str) -> float:
    total = 0.0
    wide_chars = set("MWЉЊШЖФДЏЩЮW@#%")
    narrow_chars = set("ilIјљ.,:;'|! ")
    for char in str(text or ""):
        if char.isspace():
            total += 0.35
        elif char in wide_chars:
            total += 1.25
        elif char in narrow_chars:
            total += 0.45
        elif char.isupper():
            total += 1.05
        else:
            total += 0.9
    return max(total, 1.0)


def _fit_font_size_points(paragraph, run, value: str) -> float | None:
    clean_value = " ".join(str(value or "").split())
    if len(clean_value) < 12:
        return None

    base_size = _base_font_size_points(run, paragraph)
    available_width = _available_width_points(paragraph)
    estimated_width = _weighted_character_count(clean_value) * base_size * TEXT_WIDTH_FACTOR
    if estimated_width <= available_width:
        return None

    fitted = available_width / (_weighted_character_count(clean_value) * TEXT_WIDTH_FACTOR)
    fitted = max(MIN_FIT_FONT_SIZE_PT, min(base_size, fitted))
    # Round to half-points so LibreOffice/Word render predictably.
    return round(fitted * 2) / 2


def _fit_replacement_runs(paragraph, runs, value: str) -> None:
    for run in runs:
        if not getattr(run, "text", ""):
            continue
        size = _fit_font_size_points(paragraph, run, value)
        if size is not None:
            run.font.size = Pt(size)


def _replace_in_paragraph(paragraph, placeholders: Dict[str, str]) -> None:
    if not paragraph.runs:
        return

    ordered = sorted(placeholders.items(), key=lambda kv: len(kv[0]), reverse=True)
    for key, raw_value in ordered:
        value = str(raw_value)
        while True:
            if _replace_in_single_run(paragraph, key, value):
                continue
            if _replace_across_runs(paragraph, key, value):
                continue
            break


def _replace_placeholders(doc: Document, placeholders: Dict[str, str]):
    """Replace placeholders in paragraphs and tables while preserving run styling where possible."""
    for para in _iter_all_paragraphs(doc):
        _replace_in_paragraph(para, placeholders)
